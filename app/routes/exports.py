# app/routes/exports.py
from __future__ import annotations

import io
from datetime import datetime
from flask import Blueprint, request, send_file, jsonify
from werkzeug.utils import secure_filename

# PDF / DOCX (añade a tu venv:  pip install reportlab python-docx )
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors

from docx import Document
from docx.shared import Pt, Cm

bp = Blueprint("exports", __name__)

def _norm_filename(name: str) -> str:
    base = secure_filename(name or "resultado").strip() or "resultado"
    return base

def _txt_bytes(text: str) -> bytes:
    return (text or "").encode("utf-8")

def _build_srt(transcript: str) -> str:
    """
    SRT 'sin tiempos' (placeholder). Cada línea = 3s aprox para que players no fallen.
    Si en el futuro tienes timestamps reales, sustitúyelo.
    """
    lines = [l for l in (transcript or "").splitlines() if l.strip()]
    out = []
    t = 0
    idx = 1
    step = 3
    for l in lines:
        t0 = t
        t1 = t + step
        hh0, mm0, ss0 = t0 // 3600, (t0 % 3600) // 60, t0 % 60
        hh1, mm1, ss1 = t1 // 3600, (t1 % 3600) // 60, t1 % 60
        out.append(f"{idx}")
        out.append(f"{hh0:02d}:{mm0:02d}:{ss0:02d},000 --> {hh1:02d}:{mm1:02d}:{ss1:02d},000")
        out.append(l.strip())
        out.append("")
        t += step
        idx += 1
    return "\n".join(out) or "1\n00:00:00,000 --> 00:00:02,000\n(Transcripción vacía)\n"

def _build_vtt(transcript: str) -> str:
    lines = [l for l in (transcript or "").splitlines() if l.strip()]
    out = ["WEBVTT", ""]
    t = 0
    step = 3
    for l in lines:
        t0 = t
        t1 = t + step
        hh0, mm0, ss0 = t0 // 3600, (t0 % 3600) // 60, t0 % 60
        hh1, mm1, ss1 = t1 // 3600, (t1 % 3600) // 60, t1 % 60
        out.append(f"{hh0:02d}:{mm0:02d}:{ss0:02d}.000 --> {hh1:02d}:{mm1:02d}:{ss1:02d}.000")
        out.append(l.strip())
        out.append("")
        t += step
    if len(out) == 2:
        out += ["00:00:00.000 --> 00:00:02.000", "(Transcripción vacía)"]
    return "\n".join(out)

def _build_docx(transcript: str, summary: str) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.left_margin  = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin   = Cm(2.0)
    section.bottom_margin= Cm(2.0)

    title = doc.add_paragraph()
    run = title.add_run("PolyScribe · Resultado")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph().add_run(datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")).italic = True
    doc.add_paragraph("")

    h = doc.add_paragraph()
    rh = h.add_run("Transcripción")
    rh.bold = True
    rh.font.size = Pt(13)

    for line in (transcript or "").splitlines():
        doc.add_paragraph(line)

    doc.add_paragraph("")
    h2 = doc.add_paragraph()
    rh2 = h2.add_run("Resumen")
    rh2.bold = True
    rh2.font.size = Pt(13)

    for line in (summary or "").splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def _build_pdf(transcript: str, summary: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    title = styles["Title"]
    title.textColor = colors.HexColor("#0b62e0")
    story.append(Paragraph("PolyScribe · Resultado", title))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"), styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    story.append(Paragraph("<b>Transcripción</b>", styles["Heading3"]))
    for line in (transcript or "").splitlines():
        story.append(Paragraph(line, styles["Normal"]))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("<b>Resumen</b>", styles["Heading3"]))
    for line in (summary or "").splitlines():
        story.append(Paragraph(line, styles["Normal"]))

    doc.build(story)
    buf.seek(0)
    return buf.read()

@bp.route("/api/exports/<fmt>", methods=["POST"])
def export_file(fmt: str):
    """
    POST JSON: { transcript, summary, filename }
    Devuelve attachment para SRT / VTT / DOCX / PDF.
    TXT y JSON puedes seguir haciéndolos en el front si quieres.
    """
    data = request.get_json(silent=True) or {}
    transcript = data.get("transcript", "") or ""
    summary    = data.get("summary", "") or ""
    base       = _norm_filename(data.get("filename", "resultado"))

    try:
        if fmt == "srt":
            payload = _build_srt(transcript).encode("utf-8")
            return send_file(
                io.BytesIO(payload),
                mimetype="application/x-subrip",
                as_attachment=True,
                download_name=f"{base}.srt",
            )

        if fmt == "vtt":
            payload = _build_vtt(transcript).encode("utf-8")
            return send_file(
                io.BytesIO(payload),
                mimetype="text/vtt",
                as_attachment=True,
                download_name=f"{base}.vtt",
            )

        if fmt == "docx":
            payload = _build_docx(transcript, summary)
            return send_file(
                io.BytesIO(payload),
                mimetype=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
                as_attachment=True,
                download_name=f"{base}.docx",
            )

        if fmt == "pdf":
            payload = _build_pdf(transcript, summary)
            return send_file(
                io.BytesIO(payload),
                mimetype="application/pdf",
                as_attachment=True,
                download_name=f"{base}.pdf",
            )

        return jsonify({"error": f"Formato no soportado: {fmt}"}), 400

    except Exception as e:
        return jsonify({"error": f"No se pudo exportar: {e}"}), 500
